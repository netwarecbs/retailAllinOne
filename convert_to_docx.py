#!/usr/bin/env python3
"""
Script to convert markdown documentation to DOCX format
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def setup_document_styles(doc):
    """Setup custom styles for the document"""
    
    # Title style
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(18)
    title_font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Heading 1 style
    h1_style = doc.styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Arial'
    h1_font.size = Pt(16)
    h1_font.bold = True
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)
    
    # Heading 2 style
    h2_style = doc.styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Arial'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_style.paragraph_format.space_before = Pt(10)
    h2_style.paragraph_format.space_after = Pt(4)
    
    # Heading 3 style
    h3_style = doc.styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = 'Arial'
    h3_font.size = Pt(12)
    h3_font.bold = True
    h3_style.paragraph_format.space_before = Pt(8)
    h3_style.paragraph_format.space_after = Pt(3)
    
    # Code style
    code_style = doc.styles.add_style('CodeStyle', WD_STYLE_TYPE.PARAGRAPH)
    code_font = code_style.font
    code_font.name = 'Courier New'
    code_font.size = Pt(10)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(3)
    
    # Normal paragraph style
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)

def add_code_block(doc, code_text):
    """Add a code block to the document"""
    p = doc.add_paragraph()
    p.style = 'CodeStyle'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Split code into lines and add each line
    lines = code_text.strip().split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run('\n')
        p.add_run(line)

def convert_markdown_to_docx(markdown_file, output_file):
    """Convert markdown file to DOCX format"""
    
    # Read the markdown file
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create a new document
    doc = Document()
    
    # Setup document styles
    setup_document_styles(doc)
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Handle title (first line)
        if i == 0 and line.startswith('#'):
            title = line.lstrip('#').strip()
            p = doc.add_paragraph(title)
            p.style = 'CustomTitle'
            i += 1
            continue
        
        # Handle main headings (##)
        if line.startswith('##') and not line.startswith('###'):
            heading = line.lstrip('#').strip()
            p = doc.add_heading(heading, level=1)
            i += 1
            continue
        
        # Handle subheadings (###)
        if line.startswith('###'):
            heading = line.lstrip('#').strip()
            p = doc.add_heading(heading, level=2)
            i += 1
            continue
        
        # Handle sub-subheadings (####)
        if line.startswith('####'):
            heading = line.lstrip('#').strip()
            p = doc.add_heading(heading, level=3)
            i += 1
            continue
        
        # Handle code blocks
        if line.startswith('```'):
            # Find the end of the code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                code_text = '\n'.join(code_lines)
                add_code_block(doc, code_text)
            i += 1
            continue
        
        # Handle bullet points
        if line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph(bullet_text, style='List Bullet')
            i += 1
            continue
        
        # Handle numbered lists
        if re.match(r'^\d+\.', line):
            list_text = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(list_text, style='List Number')
            i += 1
            continue
        
        # Handle table of contents
        if line.startswith('- [') and '](#' in line:
            # Extract link text
            match = re.search(r'- \[([^\]]+)\]\(#[^)]+\)', line)
            if match:
                toc_text = match.group(1)
                p = doc.add_paragraph(toc_text, style='List Bullet')
            i += 1
            continue
        
        # Handle interface definitions
        if line.startswith('```typescript') or line.startswith('```typescript'):
            # Find the end of the code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                code_text = '\n'.join(code_lines)
                add_code_block(doc, code_text)
            i += 1
            continue
        
        # Handle regular paragraphs
        if line:
            # Check if this is part of a multi-line paragraph
            paragraph_lines = [line]
            i += 1
            
            # Collect continuation lines
            while i < len(lines):
                next_line = lines[i].strip()
                if (next_line and 
                    not next_line.startswith('#') and 
                    not next_line.startswith('-') and 
                    not next_line.startswith('*') and 
                    not next_line.startswith('```') and 
                    not re.match(r'^\d+\.', next_line) and
                    not next_line.startswith('**') and
                    not next_line.startswith('✅') and
                    not next_line.startswith('**Document Version**')):
                    paragraph_lines.append(next_line)
                    i += 1
                else:
                    break
            
            # Join lines and add paragraph
            paragraph_text = ' '.join(paragraph_lines)
            
            # Handle bold text
            if paragraph_text.startswith('**') and paragraph_text.endswith('**'):
                bold_text = paragraph_text[2:-2]
                p = doc.add_paragraph()
                run = p.add_run(bold_text)
                run.bold = True
            else:
                doc.add_paragraph(paragraph_text)
        
        i += 1
    
    # Add page breaks before major sections
    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'Heading 1':
            if paragraph.text in ['System Overview', 'Inventory Stock Entry Process', 
                                'Purchase Bill Creation and Management', 'Voucher, Invoice, and Bill Linking',
                                'Data Flow Architecture', 'API Reference', 'User Interface Components',
                                'Business Logic and Validation', 'Integration Points', 'Troubleshooting Guide']:
                # Add page break before these sections
                paragraph._element.addprevious(OxmlElement('w:br'))
                paragraph._element.addprevious(OxmlElement('w:br'))
    
    # Save the document
    doc.save(output_file)
    print(f"Document converted successfully: {output_file}")

if __name__ == "__main__":
    input_file = "INVENTORY_PURCHASE_SYSTEM_DOCUMENTATION.md"
    output_file = "INVENTORY_PURCHASE_SYSTEM_DOCUMENTATION.docx"
    
    try:
        convert_markdown_to_docx(input_file, output_file)
        print("Conversion completed successfully!")
    except Exception as e:
        print(f"Error during conversion: {e}")
