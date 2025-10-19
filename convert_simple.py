#!/usr/bin/env python3
"""
Simple script to convert markdown to DOCX using python-docx
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def convert_markdown_to_docx_simple(markdown_file, output_file):
    """Simple conversion of markdown to DOCX"""
    
    # Read the markdown file
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create a new document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Handle title (first line)
        if i == 0 and line.startswith('#'):
            title = line.lstrip('#').strip()
            p = doc.add_heading(title, level=0)
            i += 1
            continue
        
        # Handle main headings (##)
        if line.startswith('##') and not line.startswith('###'):
            heading = line.lstrip('#').strip()
            doc.add_heading(heading, level=1)
            i += 1
            continue
        
        # Handle subheadings (###)
        if line.startswith('###'):
            heading = line.lstrip('#').strip()
            doc.add_heading(heading, level=2)
            i += 1
            continue
        
        # Handle sub-subheadings (####)
        if line.startswith('####'):
            heading = line.lstrip('#').strip()
            doc.add_heading(heading, level=3)
            i += 1
            continue
        
        # Handle code blocks
        if line.startswith('```'):
            # Find the end of the code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            i += 1
            continue
        
        # Handle bullet points
        if line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:].strip()
            doc.add_paragraph(bullet_text, style='List Bullet')
            i += 1
            continue
        
        # Handle numbered lists
        if re.match(r'^\d+\.', line):
            list_text = re.sub(r'^\d+\.\s*', '', line)
            doc.add_paragraph(list_text, style='List Number')
            i += 1
            continue
        
        # Handle table of contents
        if line.startswith('- [') and '](#' in line:
            # Extract link text
            match = re.search(r'- \[([^\]]+)\]\(#[^)]+\)', line)
            if match:
                toc_text = match.group(1)
                doc.add_paragraph(toc_text, style='List Bullet')
            i += 1
            continue
        
        # Handle regular paragraphs
        if line:
            # Check if this is part of a multi-line paragraph
            paragraph_lines = [line]
            i += 1
            
            # Collect continuation lines
            while i < len(lines):
                next_line = lines[i].strip()
                if (next_line and 
                    not next_line.startswith('#') and 
                    not next_line.startswith('-') and 
                    not next_line.startswith('*') and 
                    not next_line.startswith('```') and 
                    not re.match(r'^\d+\.', next_line) and
                    not next_line.startswith('**') and
                    not next_line.startswith('✅') and
                    not next_line.startswith('**Document Version**')):
                    paragraph_lines.append(next_line)
                    i += 1
                else:
                    break
            
            # Join lines and add paragraph
            paragraph_text = ' '.join(paragraph_lines)
            
            # Handle bold text
            if paragraph_text.startswith('**') and paragraph_text.endswith('**'):
                bold_text = paragraph_text[2:-2]
                p = doc.add_paragraph()
                run = p.add_run(bold_text)
                run.bold = True
            else:
                doc.add_paragraph(paragraph_text)
        
        i += 1
    
    # Save the document
    doc.save(output_file)
    print(f"Document converted successfully: {output_file}")

if __name__ == "__main__":
    input_file = "INVENTORY_PURCHASE_SYSTEM_DOCUMENTATION.md"
    output_file = "INVENTORY_PURCHASE_SYSTEM_DOCUMENTATION.docx"
    
    try:
        convert_markdown_to_docx_simple(input_file, output_file)
        print("Conversion completed successfully!")
    except Exception as e:
        print(f"Error during conversion: {e}")
        print("Make sure you have python-docx installed: pip install python-docx")
